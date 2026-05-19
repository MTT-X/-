"""Generate Word report (report.docx)"""
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)


def add_h(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')


def add_p(text, indent=True):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    run = p.add_run(text)
    run.font.size = Pt(12)
    return p


def add_tbl(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r+1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()
    return table


# ===== COVER =====
for _ in range(6):
    doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('From Text to Risk')
r.font.size = Pt(26)
r.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Shipping Insurance Intelligent Classification and Risk Assessment\nBased on Unstructured Cargo Descriptions')
r.font.size = Pt(14)
r.italic = True

for _ in range(3):
    doc.add_paragraph()

lines = [
    ('FINA2003 Financial Technology Frontiers - Final Project Report', 14),
    ('Partner Company: Anlan Information Technology', 12),
    ('Group 1', 14),
    ('Members: Hua Yu (2025241005), Mei Yuyang (2025241009), Yang Yuxuan (2025241019)', 12),
]
for text, size in lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(text).font.size = Pt(size)

for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('May 20, 2025').font.size = Pt(12)

doc.add_page_break()

# ===== MODULE 1 =====
add_h('1. Industry and Product Analysis', 1)

add_h('1.1 Shipping Insurance Market: Status and Pain Points', 2)
add_p('China\'s shipping insurance market exceeds RMB 50 billion with an annual growth rate of 8-10%, dominated by Ping An, PICC, and CPIC. In the underwriting process, cargo classification is the core of risk pricing - the cargo descriptions filled by policyholders directly affect insurers\' risk level judgments, which in turn determine whether to underwrite and at what premium.')
add_p('Core pain points: (1) Non-standardized input: policyholders fill in free-text descriptions ranging from "bulk chemical liquid" to "construction equipment parts", without unified format; (2) Heavy manual dependency: risk assessors must manually review each entry, taking 5-15 minutes per case; (3) Inconsistent classification standards: different reviewers may reach different conclusions for the same description, leading to pricing inconsistency.')
add_p('This project addresses Anlan Tech\'s real business challenge by designing and implementing an automatic classification and risk assessment solution based on unstructured cargo descriptions.')

add_h('1.2 Competitive Landscape and Solution Research', 2)
add_tbl(
    ['Solution Type', 'Representative', 'Advantages', 'Disadvantages'],
    [
        ['Rule Engine', 'In-house insurer systems', 'Controllable, interpretable', 'High maintenance, limited coverage'],
        ['ML Classification', 'Concirrus(UK), Planck(IL)', 'Adaptive, high accuracy', 'Needs labeled data, black-box'],
        ['LLM-based', 'Zest AI, etc.', 'Semantic understanding, zero-shot', 'High cost, hallucination risk'],
        ['Hybrid', 'Select MGAs', 'Rule+ML complementary', 'High complexity, cost overlap'],
    ]
)
add_p('Given constraints (no labeled data, need for interpretability, short timeline), we chose the rule-matching approach. It enables rapid deployment with high interpretability and low maintenance cost; limitations can be addressed through subsequent iterations.')

add_h('1.3 Product Design', 2)
add_p('The system uses a three-tier processing pipeline:')
add_p('(1) Text Preprocessing Layer: Clean cargo descriptions (remove newlines, fullwidth-to-halfwidth conversion, whitespace normalization).', indent=False)
add_p('(2) Exclusion Screening Layer: Keyword matching against 10 excluded cargo categories; any match triggers "High Risk (Decline)" flag.', indent=False)
add_p('(3) HS Category Matching Layer: For items not excluded, match against 22 HS major categories; assign the highest-scoring category.', indent=False)
add_p('Human Review Mechanism: Only "High Risk" and "Unclassified" cases trigger manual review (~30% of total), while the rest proceed automatically, improving efficiency by ~70%. Manual decisions on unclassified cases feed back into the keyword dictionary for continuous improvement.')

add_h('1.4 Deployment Risks and Mitigation', 2)
add_p('Classification Error Risk: Adopt conservative strategy for high-risk categories (better to over-reject than under-reject), with a human review transition period during initial deployment.')
add_p('Data Bias Risk: Establish regular dictionary update mechanisms, supplementing keywords through analysis of unclassified cases.')
add_p('Systematic Underestimation Risk: Future integration of historical claims data, overlaying loss ratio weights on top of category labels.')

# ===== MODULE 2 =====
add_h('2. Data Processing and Classification Modeling', 1)

add_h('2.1 Data Exploration', 2)
add_tbl(
    ['Dataset', 'Size', 'Content'],
    [
        ['Cargo Description 2025.xlsx', '28,304 entries', 'Free-text cargo descriptions from policyholders'],
        ['Customs HS Code Table.xlsx', '8,045 entries', 'HS code system, 22 major + detailed categories'],
        ['Insurance Exclusion List.xlsx', '10 categories', 'Goods excluded from insurance coverage'],
    ]
)
add_p('Data characteristics: ~60% English, ~40% Chinese with mixed usage; average description length ~45 characters (min 1, max 500+); ~0.9% (259 entries) are invalid descriptions like "SEE ATTACHMENT"; common noise includes container numbers, shipping info, and HS codes embedded in description text.')

add_h('2.2 Text Preprocessing', 2)
add_p('(1) Newline normalization: replace \\n, \\t and actual line breaks with spaces.')
add_p('(2) Fullwidth-to-halfwidth conversion: fullwidth characters (FF01-FF5E) to halfwidth, fullwidth space (U+3000) to halfwidth space.')
add_p('(3) Whitespace normalization: collapse consecutive spaces into single space.')
add_p('Rationale: Under the keyword-matching strategy, preprocessing aims to preserve original semantic information rather than aggressively tokenizing or removing stop words, avoiding accidental loss of classification signals.')

add_h('2.3 Classification Implementation', 2)
add_p('Method: Multi-layer rule-based keyword matching. Rationale: no training data required, fully interpretable, rules can be flexibly adjusted.')
add_p('Dictionary scale: 10 exclusion categories + 22 HS categories, ~500+ Chinese and English keywords. Chinese uses substring matching; English short words (<=3 chars) use regex word-boundary matching to avoid false positives.')
add_p('KEY FINDING - Substring False Match Problem: In the initial version, 95.3% of exclusion cases hit only 1 keyword. Investigation revealed: "car" (vehicle) falsely matched "cardboard", "art" (artwork) falsely matched "parts", "ship" (vessel) falsely matched "shipper". After fixing with word-boundary matching, the exclusion rate dropped from 28.2% to 17.5%, eliminating 3,045 false positives. This demonstrates that the real challenge for rule-based systems lies not in recall but in precision - word-boundary matching is a fundamental yet often overlooked detail in English NLP.')

add_h('2.4 Classification Results', 2)
add_p('Overview: Total 28,304 entries; Insurable 20,309 (82.5%); Excluded 4,950 (17.5%); Classification coverage 83.3%.')

add_tbl(
    ['Exclusion Category', 'Count', 'Percentage'],
    [
        ['Glass/Ceramic/Stone/Fragile', '1,470', '29.7%'],
        ['Vehicles/Aircraft/Vessels', '1,156', '23.4%'],
        ['Perishable/Fresh Goods', '630', '12.7%'],
        ['Precision Instruments/Chips/LCD', '409', '8.3%'],
        ['Flammable/Explosive', '295', '6.0%'],
    ]
)

add_tbl(
    ['HS Category (Insurable Top 5)', 'Count', 'Percentage'],
    [
        ['Machinery, Mechanical & Electrical Equipment', '3,774', '16.2%'],
        ['Base Metals & Articles', '3,462', '14.8%'],
        ['Plastics & Rubber', '2,389', '10.2%'],
        ['Chemical Industry Products', '1,756', '7.5%'],
        ['Textiles & Textile Articles', '1,552', '6.6%'],
    ]
)

add_h('2.5 Error Analysis', 2)
add_p('Remaining unclassified items: 4,718 (20.2%), mainly: (1) highly specialized product names (e.g., "PHAFFIA RHODOZYMA", "PARA ARAMID FIBER"); (2) vague descriptions (e.g., "HOUSEHOLD ITEMS"); (3) multi-product mixed descriptions. These cases are suitable for TF-IDF vector similarity matching or LLM zero-shot classification as supplementary strategies.')

# ===== MODULE 3 =====
add_h('3. Risk Prediction Exploration', 1)

add_h('3.1 Risk Labeling and Assessment', 2)
add_p('Risk stratification framework based on classification results:')
add_p('High Risk (Risk=1.0): Hits any exclusion category, 4,950 entries (17.5%) - recommend auto-decline + human review.', indent=False)
add_p('Medium-High Risk (Risk=0.5-0.9): Not directly excluded but inherently high-risk category (e.g., glass/ceramic products), ~1,800 entries (6.4%) - recommend manual pricing.', indent=False)
add_p('Low Risk (Risk=0-0.5): Conventional industrial and consumer goods, ~16,836 entries (59.5%) - recommend auto-underwrite + standard pricing.', indent=False)
add_p('Unclassified: 4,718 entries (16.7%) - requires human determination.', indent=False)

add_h('3.2 Business Implications', 2)
add_p('Approximately 59.5% of cargo can be safely auto-underwritten, and 17.5% can be automatically identified as high-risk and declined - the rule system covers ~77% of routine business scenarios, compressing manual review to ~23%. For an insurer processing hundreds of thousands of policies annually, this translates to thousands of hours saved in manual review time.')
add_p('Limitations: Current risk scoring is a preliminary inference based solely on cargo categories, not incorporating historical claims data, route risk, seasonality, etc. We recommend building a supervised learning risk prediction model (LightGBM/XGBoost) after accumulating 6-12 months of labeled data, using classification labels as one feature among many.')

# ===== SUMMARY =====
add_h('4. Summary and Outlook', 1)
add_p('Achievements: (1) Built a Chinese-English keyword dictionary covering 10 exclusion categories + 22 HS categories; (2) Implemented a three-tier rule-matching engine (~400 lines of code, ~15s runtime); (3) Achieved 83.3% automatic classification coverage; (4) Discovered and fixed the substring false-match precision issue.')
add_p('Improvement Directions: Short-term - introduce TF-IDF vector similarity as fallback; Mid-term - accumulate labeled data and train XGBoost classifier; Long-term - integrate historical claims data for end-to-end risk scorecard modeling, explore BERT-based semantic matching.')
add_p('Outlook: The "from text to risk" approach can be extended to property insurance, cargo insurance, and other insurance scenarios requiring risk assessment based on unstructured descriptions. As large language model technology matures and costs decrease, we envision a fully automated underwriting pipeline: "description input -> intelligent parsing -> risk pricing -> policy generation".')

# ===== REFERENCES =====
add_h('References', 1)
refs = [
    'China Banking and Insurance Regulatory Commission. Shipping Insurance Development Report. 2024.',
    'International Union of Marine Insurance. Global Marine Insurance Report 2024.',
    'General Administration of Customs of China. Explanatory Notes to the Customs Tariff. 2024 ed.',
    'Concirrus. AI-Powered Marine Insurance Analytics. concirrus.com',
    'Anlan Information Technology. Smart Shipping Insurance Solution White Paper. 2024.',
    'Jurafsky D, Martin J H. Speech and Language Processing. 3rd ed. 2024.',
]
for i, ref in enumerate(refs, 1):
    add_p(f'[{i}] {ref}', indent=False)

# ===== APPENDIX A =====
doc.add_page_break()
add_h('Appendix A: Code Documentation', 1)
add_p('Runtime Environment: Python 3.9+, pandas, openpyxl', indent=False)
add_p('Execution: python cargo_classifier.py (processes 28K entries in ~15 seconds)', indent=False)
add_p('Project Structure:', indent=False)
files = [
    'cargo_classifier.py - Main classification engine (~400 lines)',
    'presentation.html - Classroom presentation slides (14 slides)',
    'report.html - PDF report source (print to PDF from browser)',
    'report.docx - Word report (editable)',
    'classification_results.xlsx - Detailed results (28,304 entries)',
    'classification_report.txt - Statistical analysis report',
    'Cargo_Description_2025.xlsx - Raw data (confidential)',
    'Customs_HS_Code_Table.xlsx - Raw data (confidential)',
    'Insurance_Exclusion_List.xlsx - Raw data (confidential)',
]
for f in files:
    add_p(f'  {f}', indent=False)

# ===== APPENDIX C =====
doc.add_page_break()
add_h('Appendix C: Team Member Contributions', 1)
add_tbl(
    ['Member', 'Student ID', 'Contribution (<=150 words)'],
    [
        ['Hua Yu', '2025241005', 'Responsible for data exploration, text preprocessing, keyword dictionary construction and optimization, classification result statistical analysis, and error case diagnosis.'],
        ['Mei Yuyang', '2025241009', 'Responsible for presentation PPT creation, industry background research, competitive landscape analysis, product design proposal writing, and oral presentation delivery.'],
        ['Yang Yuxuan', '2025241019', 'Responsible for classification engine coding and debugging, rule-matching logic implementation, word-boundary fix, and system testing.'],
    ]
)
add_p('All three members participated in the writing, revision, and finalization of the PDF report.', indent=False)

# Save
output_path = r'E:\SKD\金融科技前沿\期末项目\report.docx'
doc.save(output_path)
print(f'Word report saved to: {output_path}')
